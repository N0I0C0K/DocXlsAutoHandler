import sys
import os
import xlrd
import xlwt
import time
from xlutils.copy import copy
from docx import Document,document
from mywidget import Ui_Form
from PyQt5.QtWidgets import QApplication, QWidget,QFileDialog
from PyQt5.QtCore import QEvent

class mywidget(QWidget):
    ui:Ui_Form = None
    docfilename:str = ''
    xlsfilename:str = ''
    def __init__(self, *args, **kwargs) -> None:
        super().__init__(*args, **kwargs)
        self.ui = Ui_Form()
        self.ui.setupUi(self)
        #切换功能时要重置所有选项
        self.ui.tabWidget.currentChanged.connect(self.changeFunction)
        
        #第一个模块功能
        self.ui.selectDoc.clicked.connect(self.getDocfilename)
        self.ui.selectXls.clicked.connect(self.getXlsfilename)
        self.ui.run_model_one.clicked.connect(self.run_model_one)

        #第二个模块初始化
        self.ui.selectDoc_2.clicked.connect(self.getDocfilename_model_two)
        self.ui.selectXls_2.clicked.connect(self.getXlsfilename_model_two)
        self.ui.selectDocDir.clicked.connect(self.getDocDir_model_two)
        self.ui.selectXlsDir.clicked.connect(self.getXlsDir_model_two)

        self.ui.run_model_two.clicked.connect(self.run_model_two)
        self.ui.model_two_combox.currentIndexChanged.connect(self.model_two_function_changed)
        self.ui.model_two_combox.setCurrentIndex(1)

        #第三个模块初始化
        self.ui.selectXls_3.clicked.connect(self.getXlsfilename_model_three)
        self.ui.selectXlsDir_2.clicked.connect(self.getXlsDir_model_three)
        self.ui.run_model_three.clicked.connect(self.run_model_three)

    def getDocfilename(self):
        filename, filetype = QFileDialog.getOpenFileName(self, "选取模板文件", os.getcwd(), "*.docx;*.doc;/")
        self.ui.docfilename.setText(filename)
        self.docfilename = filename
        self.logAppend(f'[*]success get doc:{filename}')
        return

    def getXlsfilename(self):
        filename, filetype = QFileDialog.getOpenFileName(self, "选取数据文件", os.getcwd(), "*.xlsx;*.xls")
        self.ui.xlsfilename.setText(filename)
        self.xlsfilename = filename
        self.logAppend(f'[*]success get xls:{filename}')
        return
    
    def run_model_one(self):
        #print(f'{self.docfilename} : {self.xlsfilename}')
        if len(self.docfilename) <= 1 or len(self.xlsfilename) <= 1:
            self.logAppend('[!]error: 请选择模板文件和数据文件')
            return
        replaceStr = 'XXX'      #默认的替换字符为XXX(大写)
        if len(self.ui.replaceStr.text()) > 0:
            replaceStr = self.ui.replaceStr.text()
        self.logAppend(f'[~]开始模板题换, 替换模板为 {self.docfilename}, 数据源为 {self.xlsfilename}, 替换关键词为{replaceStr}')
        xls = xlrd.open_workbook(self.xlsfilename)
        sheet1_content = xls.sheet_by_index(0)
        for rows in sheet1_content.get_rows():
            count = 1
            doc:document.Document = Document(self.docfilename)
            for i in doc.paragraphs:
                for run in i.runs:
                    text:str = run.text
                    if count >= sheet1_content.ncols:
                        break
                    if text.find(replaceStr) != -1:
                        run.text = text.replace(replaceStr, str(rows[count].value))
                        count+=1
            _path = self.docfilename[:self.docfilename.rfind('/')+1]
            doc.save(f'{_path+str(rows[0].value)}.docx')
            self.logAppend(f'[~]生成:{_path+str(rows[0].value)}.docx, 一共替换{count-1}处')
        self.logAppend('[*]success run complete')  
        return

    def getDocfilename_model_two(self):
        filename, filetype = QFileDialog.getOpenFileName(self, "选取模板文件", os.getcwd(), "*.docx;*.doc")
        self.ui.docfilename_2.setText(filename)
        self.docfilename = filename
        self.logAppend(f'[*]success get doc:{filename}')
        return

    def getXlsfilename_model_two(self):
        filename, filetype = QFileDialog.getOpenFileName(self, "选取数据文件", os.getcwd(), "*.xlsx;*.xls")
        self.ui.xlsfilename_2.setText(filename)
        self.xlsfilename = filename
        self.logAppend(f'[*]success get xls:{filename}')
        return

    def getXlsDir_model_two(self):
        xlsdir = QFileDialog.getExistingDirectory(self, '选取表格文件夹',os.getcwd())
        self.xlsfilename = xlsdir
        self.ui.xlsfilename_2.setText(xlsdir)
        self.logAppend(f'[*]success get xls:{xlsdir}')
        return
    
    def getDocDir_model_two(self):
        docdir = QFileDialog.getExistingDirectory(self, '选取Doc文件夹',os.getcwd())
        self.docfilename = docdir
        self.ui.docfilename_2.setText(docdir)
        self.logAppend(f'[*]success get xls:{docdir}')
        return

    def model_two_function_changed(self, index:int):
        if index == 0:
            self.ui.selectDoc_2.setEnabled(True)
            self.ui.selectDocDir.setEnabled(True)
            self.ui.selectXls_2.setEnabled(False)
            self.ui.selectXlsDir.setEnabled(False)
        elif index == 1:
            self.ui.selectDoc_2.setEnabled(False)
            self.ui.selectDocDir.setEnabled(False)
            self.ui.selectXls_2.setEnabled(True)
            self.ui.selectXlsDir.setEnabled(True)
        elif index == 2:
            self.ui.selectDoc_2.setEnabled(False)
            self.ui.selectDocDir.setEnabled(False)
            self.ui.selectXls_2.setEnabled(False)
            self.ui.selectXlsDir.setEnabled(True)
        elif index == 3:
            self.ui.selectDoc_2.setEnabled(False)
            self.ui.selectDocDir.setEnabled(False)
            self.ui.selectXls_2.setEnabled(True)
            self.ui.selectXlsDir.setEnabled(False)
        elif index == 4:
            self.ui.selectDoc_2.setEnabled(False)
            self.ui.selectDocDir.setEnabled(True)
            self.ui.selectXls_2.setEnabled(False)
            self.ui.selectXlsDir.setEnabled(False)
        self.ui.xlsfilename_2.setText('')
        self.ui.docfilename_2.setText('')
        self.docfilename = ''
        self.xlsfilename = ''

    def run_model_two(self):
        if self.docfilename == '' and self.xlsfilename == '':
            self.logAppend('[!]请选择docx文件和xlsx文件')
            return
        cur_index = self.ui.model_two_combox.currentIndex()
        if cur_index == 0:
            if self.docfilename == '':
                self.logAppend(f'[!]请选择需要导出的文件')
                return
            if self.docfilename.rfind('.doc') != -1:
                self.model_two_one(self.docfilename)
            else:
                file_list = os.listdir(self.docfilename)
                for file in file_list:
                    if file.rfind('.doc') != -1:
                        self.model_two_one(self.docfilename+'/'+file)
            pass
        elif cur_index == 1:
            if self.xlsfilename == '':
                self.logAppend(f'[!]请选择需要导出的文件')
                return
            if self.xlsfilename.rfind('.xls') != -1:
                self.model_two_two(self.xlsfilename)
            else:
                file_list = os.listdir(self.xlsfilename)
                for file in file_list:
                    if file.rfind('.xls') != -1:
                        self.model_two_two(self.xlsfilename+'/'+file)
        elif cur_index == 2:
            self.model_two_three()
        elif cur_index == 3:
            self.model_two_four()
        elif cur_index == 4:
            self.model_two_five()

    def model_two_one(self, docfile:str):
        if docfile == '':
            self.logAppend(f'[!]请选择需要导出的文件')
            return
        xls = xlwt.Workbook(encoding='utf-8')
        outfilename = self.getfilename_includepath_bypath(docfile)
        self.logAppend(f'[~]开始从 {docfile} 导出表格到 {os.getcwd()}\{outfilename}.xls')
        doc:document.Document = Document(docfile)
        count = 0
        if len(doc.tables) == 0:
            self.logAppend(f'[-]{docfile}里未找到表格')
            return
        for tab in doc.tables:
            worksheet:xlwt.Worksheet = xls.add_sheet(f'{count} tab')
            self.logAppend(f'[~]正在导出{docfile}里第{count+1}个表格')
            for i in range(len(tab.rows)):
                tar = tab.row_cells(i)
                for j in range(len(tar)):
                    worksheet.write(i,j, label=tab.cell(i,j).text)
            count += 1
        self.logAppend(f'[*]导出完毕,导出到 {outfilename}.xls')     
        xls.save(f'{outfilename}.xls')
        return

    def model_two_two(self, xlsfile:str):
        if xlsfile == '':
            self.logAppend(f'[!]请选择要导出的文件')
            return
        xls = xlrd.open_workbook(xlsfile)
        outfilename = self.docfilename if self.docfilename != '' else self.getfilename_includepath_bypath(xlsfile)+'.docx'
        doc:document.Document = Document(self.docfilename) if self.docfilename != '' else Document()
        self.logAppend(f'[~]开始从{xlsfile} 导出表格到 ')
        for item in xls.sheets():
            sheet:xlrd.sheet.Sheet = item
            self.logAppend(f'[~]开始导出 sheet:{sheet.name}')
            tabel = doc.add_table(sheet.nrows, sheet.ncols)
            tabel.style = 'Table Grid'
            for i in range(sheet.nrows):
                for j in range(sheet.ncols):
                    tabel.cell(i,j).text = str(sheet.cell(i,j).value)
        self.logAppend(f'[*]成功导出到 {outfilename}')
        doc.save(outfilename)
        return

    def model_two_three(self):
        if self.xlsfilename == '' or self.xlsfilename.rfind('.xls') != -1:
            self.logAppend(f'[!]请选择需要操作的文件夹')
            return
        file_list = os.listdir(self.xlsfilename)
        file_path = self.xlsfilename+'/'
        tar_xls = xlwt.Workbook('utf-8')
        work_sheet:xlwt.Worksheet = tar_xls.add_sheet('sheet one')
        row_now = 0
        for file in file_list:
            if file.rfind('.xls') == -1:
                continue
            self.logAppend(f'[~]开始合并{file}')
            xls = xlrd.open_workbook(file_path+file)
            for idx in range(xls.nsheets):
                sheet = xls.sheet_by_index(idx)
                for i in range(sheet.nrows):
                    for j in range(sheet.ncols):
                        work_sheet.write(row_now ,j ,sheet.cell(i,j).value)
                    row_now+=1
        self.logAppend(f'合并完成,结果生成在{file_path}out.xls')
        tar_xls.save(file_path+'out.xls')
        pass
    
    def model_two_four(self):
        if self.xlsfilename == '' or self.xlsfilename.rfind('.xls') == -1:
            self.logAppend('[!]请选择正确的文件')
            return
        
        tar_xls = xlrd.open_workbook(self.xlsfilename)
        self.logAppend(f'[~]开始处理{self.xlsfilename}')
        for i in range(tar_xls.nsheets):
            
            sheet = tar_xls.sheet_by_index(i)
            res_xls = xlwt.Workbook('utf-8')
            res_sheet:xlwt.Worksheet = res_xls.add_sheet(sheet.name)
            for i in range(sheet.nrows):
                for j in range(sheet.ncols):
                    res_sheet.write(i, j, sheet.cell(i,j).value)
            out_file_name = self.getfilename_includepath_bypath(self.xlsfilename)+'_'+sheet.name+'.xls'
            self.logAppend(f'[~]第{i}个表格处理完毕,导出在{out_file_name}')
            res_xls.save(out_file_name)
        return

    def model_two_five(self):
        file = os.listdir(self.docfilename)
        if len(file) < 2:
            self.logAppend('[!]仅支持2个文件以上简历合并')
            return
        all_key = []
        keys = []
        for i in range(2):
            if i == 0:
                doc:document.Document = Document(self.docfilename+'/'+file[i])
                if len(doc.tables) == 0:
                    self.logAppend(f'[!]error in file {file[i]}')
                    return
                for table in doc.tables:
                    for i in range(len(table.rows)):
                        tar = table.row_cells(i)
                        for j in range(len(tar)):
                            all_key.append(table.cell(i,j).text)
            else:
                doc:document.Document = Document(self.docfilename+'/'+file[i])
                if len(doc.tables) == 0:
                    self.logAppend(f'[!]error in file {file[i]}')
                    return
                for table in doc.tables:
                    for i in range(len(table.rows)):
                        tar = table.row_cells(i)
                        for j in range(len(tar)):
                            t = table.cell(i,j).text
                            if t in all_key:
                                keys.append(t)
        
        xls = xlwt.Workbook(encoding='utf-8')
        outfilename = self.docfilename+'/'+time.strftime('%y-%H-%M-%S',time.localtime(time.time()))+'.xls'
        sheet:xlwt.Worksheet = xls.add_sheet('all')
        for i in range(len(keys)):
            sheet.write(0,i, keys[i])

        r = 1
        for filename in file:
            doc:document.Document = Document(self.docfilename+'/'+filename)
            if len(doc.tables) == 0:
                self.logAppend(f'[!]error in file {file[i]}')
                return
            c = 0
            for table in doc.tables:
                for i in range(len(table.rows)):
                    tar = table.row_cells(i)
                    for j in range(len(tar)):
                        t = table.cell(i,j).text
                        if t not in keys:
                            sheet.write(r, c, t)
                            c+=1
            r += 1

        xls.save(outfilename)
        self.logAppend(f'[*]success output at {outfilename}')
        return

    def getXlsfilename_model_three(self):
        filename, filetype = QFileDialog.getOpenFileName(self, "选取数据文件", os.getcwd(), "*.xlsx;*.xls")
        self.ui.xlsfilename_3.setText(filename)
        self.xlsfilename = filename
        self.logAppend(f'[*]success get xls:{filename}')
        return

    def getXlsDir_model_three(self):
        xlsdir = QFileDialog.getExistingDirectory(self, '选取表格文件夹',os.getcwd())
        self.xlsfilename = xlsdir
        self.ui.xlsfilename_3.setText(xlsdir)
        self.logAppend(f'[*]success get xls:{xlsdir}')
        return

    def run_model_three(self):
        if self.xlsfilename == '':
            self.logAppend('[!]请选择需要处理的文件')
            return
        if self.xlsfilename.rfind('.xls') != -1:
            self.model_three_by_file(self.xlsfilename)
        else:
            file_list = os.listdir(self.xlsfilename)
            for file in file_list:
                if file.rfind('.xls') != -1:
                    self.model_three_by_file(self.xlsfilename+'/'+file)

    def model_three_by_file(self, xlsfile:str):
        if xlsfile.rfind('.xls') == -1:
            self.logAppend(f'[!]请检查{xlsfile}是否存在异样')
            return
        self.logAppend(f'[~]开始处理{xlsfile}')
        xls = xlrd.open_workbook(xlsfile)
        wxls:xlwt.Workbook = copy(xls)
        sheet = xls.sheet_by_index(0)
        work_sheet:xlwt.Worksheet = wxls.get_sheet(0)
        
        start_pos = 1
        for i in range(sheet.ncols):
            if str(sheet.cell(0,i).value) != '':
                start_pos = i
                break
        self.logAppend(f'[~]成功处理答案,答案开始于{start_pos+1}列')
        if sheet.nrows <= 2:
            self.logAppend(f'[!]没有找到需要处理的数据,进程终止,请检查 {xlsfile} 是否完好符合要求')
            return
        self.logAppend('[~]开始处理数据')
        for i in range(2, sheet.nrows):
            score = 0
            self.logAppend(f'[~]正在处理第 {i+1} 行')
            for j in range(start_pos, sheet.ncols):
                if str(sheet.cell(i,j).value) == str(sheet.cell(0,j).value):
                    score += float(sheet.cell(1, j).value)
            work_sheet.write(i, sheet.ncols, score)
        outfilename = self.getfilename_includepath_bypath(xlsfile)+'_out.xls'
        wxls.save(outfilename)
        self.logAppend(f'处理完毕,文件导出在{outfilename}')

    def changeFunction(self, index):
        #切换功能后需要清空所有之前存储的文件路径
        self.docfilename = ''
        self.xlsfilename = ''
        self.ui.docfilename.setText('')
        self.ui.docfilename_2.setText('')
        self.ui.xlsfilename.setText('')
        self.ui.xlsfilename_2.setText('')
        self.ui.xlsfilename_3.setText('')
        self.logAppend(f'[*]success change funtion to :{self.ui.tabWidget.tabText(index)}')
        return

    def logAppend(self, text:str):
        self.ui.outputText_model_one.append(text)
    
    def getfilenamebypath(self, path:str) -> str:
        res = path[path.rfind('/')+1: path.rfind('.')]
        return res

    def getfilename_includepath_bypath(self, path:str) -> str:
        res = path[:path.rfind('.')]
        return res

def main():
    app = QApplication(sys.argv)
    wid = mywidget()
    wid.show()
    sys.exit(app.exec_())
    pass

if __name__ == '__main__':
    main()